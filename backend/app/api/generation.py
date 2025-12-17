import uuid
import io
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.api.deps import get_db, get_current_user
from app.models import User, Document
from app.services.document_generator import DocumentGenerator

router = APIRouter()


@router.post("/documents/{document_id}/generate")
def generate_document(
    document_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate content for all sections in a document."""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )

    if not document.sections:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections to generate",
        )

    generator = DocumentGenerator(db)
    results = generator.generate_document(str(document_id))

    return {
        "document_id": str(document_id),
        "status": "completed",
        "results": results,
    }


@router.post("/documents/{document_id}/sections/{section_id}/generate")
def regenerate_section(
    document_id: uuid.UUID,
    section_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Regenerate content for a specific section."""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )

    generator = DocumentGenerator(db)
    result = generator.regenerate_section(str(document_id), str(section_id))

    return result


@router.get("/documents/{document_id}/export")
def export_document(
    document_id: uuid.UUID,
    format: str = "markdown",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Export document to specified format."""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id,
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )

    # Build markdown content
    markdown_content = f"# {document.title}\n\n"

    for section in document.sections:
        if not section.is_included:
            continue

        markdown_content += f"## {section.title}\n\n"

        if section.generated_content:
            markdown_content += section.generated_content[0].content
        else:
            markdown_content += "_No content generated yet._"

        markdown_content += "\n\n"

    if format == "markdown":
        return StreamingResponse(
            io.BytesIO(markdown_content.encode('utf-8')),
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename={document.title}.md"},
        )

    elif format == "docx":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt
            import markdown

            doc = DocxDocument()
            doc.add_heading(document.title, 0)

            for section in document.sections:
                if not section.is_included:
                    continue

                doc.add_heading(section.title, level=1)

                if section.generated_content:
                    content = section.generated_content[0].content
                    # Simple paragraph addition (proper markdown conversion would be more complex)
                    for para in content.split('\n\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())

            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)

            return StreamingResponse(
                docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={document.title}.docx"},
            )

        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="DOCX export not available",
            )

    elif format == "pdf":
        try:
            from fpdf import FPDF
            import re

            def sanitize_for_pdf(text: str) -> str:
                """Remove or replace characters not supported by basic PDF fonts."""
                # Replace box drawing and other special characters
                replacements = {
                    '├': '|--', '└': '`--', '│': '|', '─': '-',
                    '┌': '+--', '┐': '--+', '┘': '--+', '┤': '|',
                    '•': '*', '→': '->', '←': '<-', '↑': '^', '↓': 'v',
                    '✓': '[x]', '✗': '[ ]', '★': '*', '☆': '*',
                    '"': '"', '"': '"', ''': "'", ''': "'",
                    '…': '...', '–': '-', '—': '-',
                }
                for old, new in replacements.items():
                    text = text.replace(old, new)
                # Remove any remaining non-Latin1 characters
                text = text.encode('latin-1', errors='replace').decode('latin-1')
                return text

            class PDF(FPDF):
                def header(self):
                    self.set_font('Helvetica', 'B', 12)
                    self.cell(0, 10, sanitize_for_pdf(document.title), align='C', new_x='LMARGIN', new_y='NEXT')
                    self.ln(5)

                def footer(self):
                    self.set_y(-15)
                    self.set_font('Helvetica', 'I', 8)
                    self.cell(0, 10, f'Page {self.page_no()}', align='C')

            pdf = PDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Title
            pdf.set_font('Helvetica', 'B', 24)
            pdf.cell(0, 15, sanitize_for_pdf(document.title), new_x='LMARGIN', new_y='NEXT')
            pdf.ln(10)

            for section in document.sections:
                if not section.is_included:
                    continue

                # Section title
                pdf.set_font('Helvetica', 'B', 16)
                pdf.set_text_color(51, 51, 51)
                pdf.cell(0, 10, sanitize_for_pdf(section.title), new_x='LMARGIN', new_y='NEXT')
                pdf.ln(3)

                # Section content
                if section.generated_content:
                    content = section.generated_content[0].content
                    # Strip markdown formatting for PDF
                    content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)  # Bold
                    content = re.sub(r'\*(.+?)\*', r'\1', content)  # Italic
                    content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)  # Headers
                    content = re.sub(r'`(.+?)`', r'\1', content)  # Inline code
                    content = re.sub(r'```[\s\S]*?```', '[Code Block]', content)  # Code blocks
                    content = sanitize_for_pdf(content)

                    pdf.set_font('Helvetica', '', 11)
                    pdf.set_text_color(0, 0, 0)
                    pdf.multi_cell(0, 6, content)
                else:
                    pdf.set_font('Helvetica', 'I', 11)
                    pdf.set_text_color(128, 128, 128)
                    pdf.cell(0, 10, 'No content generated yet.', new_x='LMARGIN', new_y='NEXT')

                pdf.ln(8)

            # Output PDF
            pdf_bytes = io.BytesIO()
            pdf.output(pdf_bytes)
            pdf_bytes.seek(0)

            return StreamingResponse(
                pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={document.title}.pdf"},
            )

        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="PDF export not available - fpdf2 not installed",
            )

    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported format: {format}. Use markdown, docx, or pdf.",
        )
